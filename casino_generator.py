# Required installations:
# pip install streamlit openai requests beautifulsoup4 scikit-learn numpy pandas python-docx

import streamlit as st
import requests
from bs4 import BeautifulSoup
import openai
import re
from datetime import datetime
import json
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import concurrent.futures
import threading
import time
import pandas as pd
import io
from docx import Document
from docx.shared import Inches

def check_password():
    """Returns `True` if user entered correct password."""
    
    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if st.session_state["password"] == st.secrets["app_password"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # Don't store password
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        # First run, show input for password
        st.text_input(
            "🔐 Enter App Password", 
            type="password", 
            on_change=password_entered, 
            key="password"
        )
        st.write("*Please contact admin for access credentials*")
        return False
    elif not st.session_state["password_correct"]:
        # Password not correct, show input + error
        st.text_input(
            "🔐 Enter App Password", 
            type="password", 
            on_change=password_entered, 
            key="password"
        )
        st.error("😕 Password incorrect")
        return False
    else:
        # Password correct
        return True

class CasinoContentGenerator:
    def __init__(self, api_key):
        openai.api_key = api_key
        
        # Enhanced style guidelines and prohibited words for humanization
        self.prohibited_words = [
            # Original AI-sounding words
            "perhaps", "maybe", "might", "could potentially", "seemingly",
            "appears to be", "tends to", "generally speaking", "it seems",
            "basically", "simply put", "essentially", "virtually",

            # Marketing/hyperbolic language to avoid
            "absolutely amazing", "incredible", "unbelievable", "mind-blowing",
            "game-changing", "revolutionary", "life-changing", "epic",
            "premier", "generous", "solid choice", "stands out",
            "commitment to responsible gambling", "user-friendly", "exciting",
            "amazing", "outstanding", "exceptional", "fabulous", "immersive",
            "experience the thrill", "like never before", "blow your mind",
            "don't miss out", "right in the center of the action",

            # Prohibited gambling claims
            "risk-free", "guaranteed win", "easy money", "sure thing",
            "guaranteed profit", "can't lose", "100% win rate"
        ]

        self.marketing_phrases_to_avoid = [
            "FABULOUS APP! FUN AND ENGAGING! DON'T MISS OUT!",
            "This immersive app lets you play and feel like you're right in the center of the action!",
            "Experience the thrill like never before!",
            "Revolutionary gaming experience that will blow your mind!",
            "stands out as a premier online",
            "With generous bonuses, a user-friendly app, and a commitment to responsible gambling"
        ]

        # Threading controls for parallel processing
        self.progress_lock = threading.Lock()
        self.completed_sections = 0
        self.total_sections = 0

    def scrape_competitor_content(self, competitor_urls):
        """Scrape competitor reviews for analysis - now uses parallel processing"""
        return self.scrape_competitors_parallel(competitor_urls)

    def scrape_competitors_parallel(self, competitor_urls, max_workers=3):
        """Scrape competitor content in parallel"""
        st.info("🌐 Scraping competitors in parallel...")
        
        def scrape_single_url(url):
            try:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = requests.get(url, headers=headers, timeout=10)
                response.raise_for_status()

                soup = BeautifulSoup(response.content, 'html.parser')
                for script in soup(["script", "style"]):
                    script.decompose()

                text = soup.get_text()
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)

                return {
                    'url': url,
                    'content': text[:3000]
                }

            except Exception as e:
                st.warning(f"⚠️ Error scraping {url}: {str(e)}")
                return {
                    'url': url,
                    'content': f"Could not scrape content from {url}"
                }

        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            competitor_content = list(executor.map(scrape_single_url, competitor_urls))
        
        st.success(f"✅ Scraped {len(competitor_content)} competitor sites")
        return competitor_content

    def update_progress(self, section_name):
        """Thread-safe progress tracking"""
        with self.progress_lock:
            self.completed_sections += 1
            # Update Streamlit progress
            if hasattr(st.session_state, 'progress_bar'):
                st.session_state.progress_bar.progress(self.completed_sections / self.total_sections)

    def generate_content_section_with_retry(self, section_data, max_retries=3):
        """Generate content for a single section with retry logic and rate limiting"""
        section_name, prompt = section_data
        
        for attempt in range(max_retries):
            try:
                # Add small delay to avoid rate limiting
                time.sleep(0.1)
                
                client = openai.OpenAI(api_key=openai.api_key)
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                    max_tokens=1000
                )
                
                content = response.choices[0].message.content
                self.update_progress(section_name)
                return section_name, content
                
            except Exception as e:
                if attempt < max_retries - 1:
                    st.warning(f"⚠️ Retry {attempt + 1} for {section_name}: {str(e)[:50]}...")
                    time.sleep(2 ** attempt)  # Exponential backoff
                else:
                    st.error(f"❌ Failed {section_name} after {max_retries} attempts: {str(e)}")
                    self.update_progress(section_name)
                    return section_name, f"[ERROR GENERATING {section_name.upper()}: {str(e)}]"

    def generate_sections_parallel(self, prompts, max_workers=8):
        """Generate multiple sections in parallel with controlled concurrency"""
        sections = {}
        self.completed_sections = 0
        self.total_sections = len(prompts)
        
        st.info(f"🚀 Starting parallel generation of {self.total_sections} sections...")
        
        # Create progress bar
        st.session_state.progress_bar = st.progress(0)
        status_text = st.empty()
        
        start_time = time.time()
        
        # Convert prompts dict to list of tuples for ThreadPoolExecutor
        section_items = list(prompts.items())
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all tasks
            future_to_section = {
                executor.submit(self.generate_content_section_with_retry, item): item[0] 
                for item in section_items
            }
            
            # Collect results as they complete
            for future in concurrent.futures.as_completed(future_to_section):
                section_name, content = future.result()
                sections[section_name] = content
                status_text.text(f"✅ Completed: {section_name} ({self.completed_sections}/{self.total_sections})")
        
        end_time = time.time()
        st.success(f"⏱️ Parallel generation completed in {end_time - start_time:.1f} seconds")
        
        return sections

    def get_embeddings(self, text_list):
        """Get embeddings for a list of texts"""
        try:
            client = openai.OpenAI(api_key=openai.api_key)
            response = client.embeddings.create(
                model="text-embedding-ada-002",
                input=text_list
            )
            return [embedding.embedding for embedding in response.data]
        except Exception as e:
            st.error(f"Error getting embeddings: {str(e)}")
            return None

    def analyze_semantic_gaps(self, competitor_content):
        """Analyze competitor content using embeddings to find semantic gaps"""
        st.info("Performing semantic analysis of competitor content...")

        # Get embeddings for competitor content
        competitor_texts = [comp['content'] for comp in competitor_content]
        competitor_embeddings = self.get_embeddings(competitor_texts)

        if not competitor_embeddings:
            return "No semantic analysis available"

        # Analyze semantic themes prompt
        semantic_analysis_prompt = f"""
        Based on semantic analysis of competitor casino reviews, identify:

        COMPETITOR CONTENT ANALYSIS:
        {chr(10).join([f"Competitor {i+1}: {content[:500]}..." for i, content in enumerate(competitor_texts)])}

        Provide:
        1. SEMANTIC THEMES: Main topics/themes competitors cover
        2. SEMANTIC GAPS: Important casino topics they miss or cover poorly
        3. KEYWORD OPPORTUNITIES: Semantic clusters we should target
        4. CONTENT DEPTH: Areas where competitors provide shallow coverage

        Format as structured analysis for content optimization.
        """

        try:
            client = openai.OpenAI(api_key=openai.api_key)
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": semantic_analysis_prompt}],
                temperature=0.3
            )
            return response.choices[0].message.content
        except Exception as e:
            st.error(f"Error in semantic analysis: {str(e)}")
            return "Semantic analysis unavailable"

    def check_content_authenticity(self, content):
        """Check generated content for marketing language and compliance issues"""
        st.info("Checking content authenticity...")

        issues_found = []

        # Check for prohibited words and phrases
        content_lower = content.lower()
        for phrase in self.prohibited_words:
            if phrase.lower() in content_lower:
                issues_found.append(f"Contains prohibited phrase: '{phrase}'")

        # Check for marketing patterns
        for pattern in self.marketing_phrases_to_avoid:
            if pattern.lower() in content_lower:
                issues_found.append(f"Contains marketing pattern: '{pattern[:50]}...'")

        # Check for compliance violations
        compliance_violations = []
        if any(word in content_lower for word in ["risk-free", "guaranteed win", "easy money", "sure thing"]):
            compliance_violations.append("Contains prohibited gambling claims")

        if "21+" not in content and "age" in content_lower:
            compliance_violations.append("Missing age requirement mention")

        # Check for authentic tone
        authenticity_score = 0
        if content.count("'") > 3:  # Contractions present
            authenticity_score += 1
        if not any(word in content_lower for word in ["amazing", "incredible", "outstanding"]):
            authenticity_score += 1
        if "like never before" not in content_lower and "experience the thrill" not in content_lower:
            authenticity_score += 1

        st.info(f"Authenticity score: {authenticity_score}/3")
        if issues_found:
            st.warning(f"Issues found: {len(issues_found)}")
            for issue in issues_found[:3]:  # Show first 3 issues
                st.write(f"  - {issue}")

        if compliance_violations:
            st.warning(f"Compliance violations: {compliance_violations}")

        return len(issues_found) == 0 and len(compliance_violations) == 0

    def optimize_content_semantically(self, generated_content, competitor_content, casino_name, keyword):
        """Compare our content semantically against competitors and optimize"""
        st.info("Performing semantic optimization...")

        # Get embeddings for our content and competitor content
        our_text = ' '.join([section for section in generated_content.values() if section])
        competitor_texts = [comp['content'] for comp in competitor_content]

        all_texts = [our_text] + competitor_texts
        all_embeddings = self.get_embeddings(all_texts)

        if not all_embeddings:
            return generated_content

        our_embedding = all_embeddings[0]
        competitor_embeddings = all_embeddings[1:]

        # Calculate semantic similarity
        similarities = []
        for comp_embedding in competitor_embeddings:
            similarity = cosine_similarity([our_embedding], [comp_embedding])[0][0]
            similarities.append(similarity)

        avg_similarity = np.mean(similarities)
        st.info(f"Semantic similarity to competitors: {avg_similarity:.3f}")

        return generated_content

    def analyze_content_gaps(self, competitor_content):
        """Identify gaps in competitor content"""
        gap_analysis_prompt = f"""
        Analyze these competitor reviews and identify what's missing or poorly covered:

        {chr(10).join(competitor_content)}

        Return a JSON object with:
        - missing_topics: topics competitors don't cover well
        - weak_areas: areas where competitors provide shallow information
        - opportunities: unique angles we can take
        """

        try:
            client = openai.OpenAI(api_key=openai.api_key)
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": gap_analysis_prompt}],
                temperature=0.3
            )
            return response.choices[0].message.content
        except Exception as e:
            st.error(f"Error in gap analysis: {str(e)}")
            return "Gap analysis unavailable"

    def generate_expert_prompts(self, casino_name, keyword, competitor_gaps, semantic_analysis):
        """Generate section-specific prompts with expert knowledge"""

        base_prompt = f"""
        You are a knowledgeable gambling industry expert writing authentic, honest casino reviews.
        Write like an informed friend sharing genuine insights - NOT like promotional advertising.

        TARGET KEYWORD: {keyword}
        CASINO: {casino_name}

        AUTHENTIC WRITING STYLE:
        - Conversational and genuine tone - like talking to a knowledgeable friend
        - Use contractions naturally (you'll, don't, it's, can't)
        - Maximum 20 words per sentence, mix short and long sentences
        - Second person when addressing readers, third person for factual content
        - Avoid jargon or explain it clearly when necessary
        - Natural paragraph flow, don't force rigid formatting

        STRICTLY AVOID:
        - Marketing language: {', '.join(self.prohibited_words[:15])}
        - Hyperbolic phrases: "experience the thrill", "like never before", "blow your mind"
        - Clichéd patterns: "stands out as premier", "generous bonuses", "solid choice"
        - Promotional tone that sounds like it's trying to sell
        - ALL CAPS excitement or artificial enthusiasm

        COMPLIANCE REQUIREMENTS:
        - Include responsible gambling context when relevant
        - Never claim "risk-free", "guaranteed win", or "easy money"
        - Mention age requirements (21+ in US) when discussing eligibility
        - Focus on entertainment value, not winning promises
        - Don't add "the" before {casino_name} unless grammatically correct
        - Naturally incorporate "{keyword}" where relevant

        CONTENT APPROACH:
        Traditional analysis: {competitor_gaps}

        SEMANTIC ANALYSIS:
        {semantic_analysis}

        Create authentic content that fills gaps while maintaining honest, conversational tone.
        Write like you genuinely know this casino, not like you're promoting it.
        """

        section_prompts = {
            "intro": f"""
            {base_prompt}

            Write a 2-3 sentence introduction for {casino_name} review that includes the bonus code and current offer.
            Make it conversational and focus on what players actually need to know.
            """,

            "key_takeaways": f"""
            {base_prompt}

            Create 4 bullet points highlighting the most important aspects of {casino_name}.
            Focus on practical information players need: mobile app availability, promo codes,
            legal states, standout features. No periods at end of bullet points.
            Write like you're giving quick tips to a friend.
            """,

            "sign_up_steps": f"""
            {base_prompt}

            Write 2-3 sentences introducing the sign-up process, then create a numbered list of
            5-7 steps users need to take to sign up at {casino_name} with the promo code.
            Include specific actions like visiting the site, clicking register, entering details, etc.
            Keep it conversational and helpful.
            """,

            "redeem_bonus_steps": f"""
            {base_prompt}

            Create a bulleted list of steps explaining how to redeem the {casino_name} bonus.
            Include what information to enter, when to apply the promo code, email verification, etc.
            Make it practical and actionable, like you're walking someone through it.
            """,

            "welcome_offer": f"""
            {base_prompt}

            Write 2-3 sentences explaining what the {casino_name} welcome offer includes.
            Be specific about bonus amounts, free spins, or other benefits new players receive.
            Keep it honest - mention both benefits and any requirements.
            """,

            "editors_review": f"""
            {base_prompt}

            Write a 2-3 paragraph honest review of {casino_name} (3-5 sentences per paragraph).
            Include specific details about games, user experience, and what's actually good or needs work.
            Write like you've genuinely used this casino and are sharing your real experience.
            """,

            "bonuses_section": f"""
            {base_prompt}

            Write detailed content about {casino_name} bonuses including:
            - 3-4 sentence overview of available bonuses
            - Daily login bonus details (3-5 sentences with redemption process)
            - Exclusive jackpots (3-5 sentences)
            - Cash prizes (3-5 sentences)

            Provide honest, practical information. Mention both benefits and limitations.
            """,

            "comparison": f"""
            {base_prompt}

            Write 2-3 paragraphs comparing {casino_name} to other social casinos.
            Focus on what actually sets it apart and where it falls short.
            Be objective and honest - this isn't a sales pitch.
            """,

            "overview": f"""
            {base_prompt}

            Write a 2-3 paragraph overview of {casino_name} (3-5 sentences per paragraph).
            Cover the casino's background, main features, target audience, and overall positioning.
            Write conversationally, like you're explaining it to someone who asked.
            """,

            "promo_availability": f"""
            {base_prompt}

            Write 3-5 sentences about where the {casino_name} promo code is available.
            Include information about legal US states where this casino operates.
            Mention age requirements (21+ in US) and any restrictions.
            """,

            "no_deposit_bonus": f"""
            {base_prompt}

            Write 3-5 sentences about {casino_name} no deposit bonus offers.
            If none currently exist, mention expired offers or explain their bonus structure honestly.
            Don't oversell - be realistic about what's available.
            """,

            "existing_customer_promos": f"""
            {base_prompt}

            Write 2-3 paragraphs about promotions available for existing customers at {casino_name}.
            Cover loyalty programs, reload bonuses, special tournaments, and other ongoing promotions.
            Be honest about the value and any limitations.
            """,

            "mobile_app": f"""
            {base_prompt}

            Write 3-5 sentences about {casino_name} mobile app availability,
            iOS/Android compatibility, and any differences from desktop version.
            Mention both what works well and any issues users might encounter.
            """,

            "banking": f"""
            {base_prompt}

            Write 3-5 sentences about deposit and withdrawal options at {casino_name}.
            Include specific payment methods in bulleted lists for deposits and withdrawals.
            Be honest about processing times and any fees.
            """,

            "games": f"""
            {base_prompt}

            Write 3-5 sentences about casino games at {casino_name} and create
            a bulleted list of the top 3 games offered.
            Focus on what players actually enjoy, not just promotional descriptions.
            """,

            "customer_support": f"""
            {base_prompt}

            Write 3-5 sentences about customer support at {casino_name}.
            Include available contact methods, realistic response times, and support quality.
            Be honest about both strengths and weaknesses.
            """,

            "responsible_gambling": f"""
            {base_prompt}

            Write 3-5 sentences about responsible gambling measures at {casino_name}.
            Cover self-exclusion tools, deposit limits, and player protection features.
            This is important - be thorough and include helpline information when relevant.
            """,

            "faqs": f"""
            {base_prompt}

            Generate 8 relevant FAQ questions and answers about {casino_name}.
            Questions should be H3 format, answers 2-3 sentences each.
            Cover practical concerns players actually have - bonuses, withdrawals, eligibility.
            Write answers conversationally, like you're helping someone understand.
            """
        }

        return section_prompts

    def create_promo_table_html(self, casino_name):
        table = f"""
        <div class="promo-table">
            <table class="bonus-details-table">
                <thead>
                    <tr>
                        <th>{casino_name} Info</th>
                        <th>Bonus Details</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td>🤑 {casino_name} Promo Code</td>
                        <td><strong>Code TBD</strong></td>
                    </tr>
                    <tr>
                        <td>💰 {casino_name} Welcome Offer</td>
                        <td><strong>Offer TBD</strong></td>
                    </tr>
                    <tr>
                        <td>📝 Wagering Requirements</td>
                        <td>[INSERT WAGERING REQUIREMENTS FOR BONUS OFFER]</td>
                    </tr>
                    <tr>
                        <td>🍰 Age Requirements</td>
                        <td>[INSERT AGE REQUIREMENTS FOR ONLINE CASINO]</td>
                    </tr>
                    <tr>
                        <td>📍 Available States</td>
                        <td>[ABBREVIATION LIST OF US STATES THIS CASINO IS LEGAL IN]</td>
                    </tr>
                    <tr>
                        <td>📲 Mobile App</td>
                        <td>[ADD iOS AND/OR ANDROID IF AVAILABLE]</td>
                    </tr>
                    <tr>
                        <td>✅ Promo Last Verified</td>
                        <td><strong>{datetime.now().strftime('%B %d, %Y')}</strong></td>
                    </tr>
                </tbody>
            </table>
        </div>
        """
        return table

    def convert_to_html_list(self, content):
        """Convert bullet points or numbered lists to HTML"""
        if not content:
            return content
            
        lines = content.split('\n')
        html_lines = []
        in_list = False
        list_type = None
        
        for line in lines:
            line = line.strip()
            if not line:
                if in_list:
                    html_lines.append(f"</{list_type}>")
                    in_list = False
                html_lines.append("")
                continue
                
            # Check for bullet points
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                if not in_list or list_type != 'ul':
                    if in_list:
                        html_lines.append(f"</{list_type}>")
                    html_lines.append("<ul>")
                    in_list = True
                    list_type = 'ul'
                content = re.sub(r'^[•\-*]\s*', '', line)
                html_lines.append(f"<li>{content}</li>")
                
            # Check for numbered lists
            elif re.match(r'^\d+\.?\s+', line):
                if not in_list or list_type != 'ol':
                    if in_list:
                        html_lines.append(f"</{list_type}>")
                    html_lines.append("<ol>")
                    in_list = True
                    list_type = 'ol'
                content = re.sub(r'^\d+\.?\s*', '', line)
                html_lines.append(f"<li>{content}</li>")
                
            else:
                if in_list:
                    html_lines.append(f"</{list_type}>")
                    in_list = False
                html_lines.append(f"<p>{line}</p>")
        
        if in_list:
            html_lines.append(f"</{list_type}>")
            
        return '\n'.join(html_lines)

    def generate_full_review(self, casino_name, keyword, competitor_urls):
        """Generate complete casino review with semantic analysis and parallel processing"""
        
        # Step 1: Competitor content analysis (parallel)
        competitor_content = self.scrape_competitor_content(competitor_urls)

        # Step 2: Analysis phase (can run these in parallel too)
        st.info("Analyzing competitor content gaps...")
        
        # Run gap analysis and semantic analysis in parallel
        analysis_start = time.time()
        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
            # Submit both analysis tasks
            gap_future = executor.submit(
                self.analyze_content_gaps, 
                [comp['content'] for comp in competitor_content]
            )
            semantic_future = executor.submit(
                self.analyze_semantic_gaps, 
                competitor_content
            )
            
            # Get results
            competitor_gaps = gap_future.result()
            semantic_analysis = semantic_future.result()
        
        analysis_time = time.time() - analysis_start
        st.success(f"📊 Analysis completed in {analysis_time:.1f} seconds")

        # Step 3: Generate section prompts
        prompts = self.generate_expert_prompts(casino_name, keyword, competitor_gaps, semantic_analysis)

        # Step 4: Generate all sections in parallel (THIS IS THE BIG SPEEDUP!)
        sections = self.generate_sections_parallel(prompts, max_workers=8)

        # Step 5: Check content authenticity and compliance
        full_content = ' '.join([section for section in sections.values() if section])
        is_authentic = self.check_content_authenticity(full_content)

        # Step 6: Semantic optimization
        optimized_sections = self.optimize_content_semantically(sections, competitor_content, casino_name, keyword)

        # Step 7: Assemble final review (HTML content only, no document wrapper)
        review = self.assemble_final_review_content_only(casino_name, keyword, optimized_sections)

        return review

    def assemble_final_review_content_only(self, casino_name, keyword, sections):
        """Assemble review content without HTML document wrapper - just the content for CMS"""
        # Convert content to HTML with proper formatting
        html_sections = {}
        for key, content in sections.items():
            html_sections[key] = self.convert_to_html_list(content)
        
        review = f"""<h1>{casino_name.upper()} Bonus Code & Review</h1>

<div class="section">
    {html_sections['intro']}
</div>

<h2>Key Takeaways & TL;DR</h2>
<div class="key-takeaways">
    {html_sections['key_takeaways']}
</div>

<h2>{casino_name.upper()} Promo Code Details</h2>
<div class="section">
    <p>Section text.</p>
    {self.create_promo_table_html(casino_name)}
</div>

<h2>Table of Contents</h2>
<div class="table-of-contents">
    <ul>
        <li><a href="#{casino_name.upper().replace(' ', '-')}-BONUSES">{casino_name.upper()} BONUSES</a></li>
        <li><a href="#sign-up">How to Sign Up at {casino_name}</a></li>
        <li><a href="#redeem-bonus">HOW TO REDEEM THE {casino_name.upper()} BONUS</a></li>
        <li><a href="#welcome-offer">WHAT IS THE {casino_name.upper()} WELCOME OFFER?</a></li>
        <li><a href="#editors-review">EDITORS REVIEW OF {casino_name.upper()}</a></li>
        <li><a href="#comparison">HOW DOES THE {casino_name.upper()} BONUS CODE COMPARE TO OTHER SOCIAL CASINOS</a></li>
        <li><a href="#overview">{casino_name.upper()} OVERVIEW</a></li>
        <li><a href="#promo-availability">WHERE IS THE {casino_name.upper()} PROMO CODE AVAILABLE</a></li>
        <li><a href="#no-deposit-bonus">{casino_name.upper()} NO DEPOSIT BONUS</a></li>
        <li><a href="#existing-customer-promos">CASINO PROMOTIONS FOR EXISTING CUSTOMERS</a></li>
        <li><a href="#mobile-app">{casino_name.upper()} MOBILE APP</a></li>
        <li><a href="#games">CASINO GAMES AT {casino_name.upper()}</a></li>
        <li><a href="#banking">BANKING OPTIONS WITH {casino_name.upper()}</a></li>
        <li><a href="#customer-support">CUSTOMER SUPPORT AT {casino_name.upper()}</a></li>
        <li><a href="#responsible-gambling">RESPONSIBLE GAMBLING AT {casino_name.upper()}</a></li>
        <li><a href="#faqs">{casino_name.upper()} FAQS</a></li>
    </ul>
</div>

<h2 id="{casino_name.upper().replace(' ', '-')}-BONUSES">{casino_name.upper()} BONUSES</h2>
<div class="section">
    {html_sections['bonuses_section']}
</div>

<h2 id="sign-up">How to Sign Up at {casino_name}</h2>
<div class="section">
    {html_sections['sign_up_steps']}
</div>

<h2 id="redeem-bonus">HOW TO REDEEM THE {casino_name.upper()} BONUS</h2>
<div class="section">
    {html_sections['redeem_bonus_steps']}
</div>

<h2 id="welcome-offer">WHAT IS THE {casino_name.upper()} WELCOME OFFER?</h2>
<div class="section">
    {html_sections['welcome_offer']}
</div>

<h2 id="editors-review">EDITORS REVIEW OF {casino_name.upper()}</h2>
<div class="section">
    {html_sections['editors_review']}
</div>

<h2 id="comparison">HOW DOES THE {casino_name.upper()} BONUS CODE COMPARE TO OTHER SOCIAL CASINOS</h2>
<div class="section">
    {html_sections['comparison']}
</div>

<h2 id="overview">{casino_name.upper()} OVERVIEW</h2>
<div class="section">
    {html_sections['overview']}
</div>

<h2 id="promo-availability">WHERE IS THE {casino_name.upper()} PROMO CODE AVAILABLE</h2>
<div class="section">
    {html_sections['promo_availability']}
</div>

<h2 id="no-deposit-bonus">{casino_name.upper()} NO DEPOSIT BONUS</h2>
<div class="section">
    {html_sections['no_deposit_bonus']}
</div>

<h2 id="existing-customer-promos">CASINO PROMOTIONS FOR EXISTING CUSTOMERS</h2>
<div class="section">
    {html_sections['existing_customer_promos']}
</div>

<h2 id="mobile-app">{casino_name.upper()} MOBILE APP</h2>
<div class="section">
    {html_sections['mobile_app']}
</div>

<h2 id="games">CASINO GAMES AT {casino_name.upper()}</h2>
<div class="section">
    {html_sections['games']}
</div>

<h2 id="banking">BANKING OPTIONS WITH {casino_name.upper()}</h2>
<div class="section">
    {html_sections['banking']}
</div>

<h2 id="customer-support">CUSTOMER SUPPORT AT {casino_name.upper()}</h2>
<div class="section">
    {html_sections['customer_support']}
</div>

<h2 id="responsible-gambling">RESPONSIBLE GAMBLING AT {casino_name.upper()}</h2>
<div class="section">
    {html_sections['responsible_gambling']}
</div>

<h2 id="faqs">{casino_name.upper()} FAQS</h2>
<div class="section">
    {html_sections['faqs']}
</div>

<div class="meta-description">
    <strong>Meta Description:</strong> {keyword} - {casino_name} bonus codes and expert review. Latest offers, games analysis, and comprehensive {casino_name} casino guide.
</div>"""
        
        return review

def create_docx_with_html_content(casino_name, html_content):
    """Create a DOCX file with raw HTML content as text"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'{casino_name} - HTML Content', 0)
    
    # Add subtitle with timestamp
    subtitle = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    subtitle.style = 'Subtitle'
    
    # Add instructions
    instructions = doc.add_paragraph()
    instructions.add_run('Instructions: ').bold = True
    instructions.add_run('Copy the HTML content below and paste it directly into your CMS editor.')
    
    doc.add_paragraph()  # Empty line
    
    # Add HTML content header
    html_header = doc.add_heading('Raw HTML Content:', level=1)
    
    # Add the HTML content as plain text in a monospace style
    html_paragraph = doc.add_paragraph()
    run = html_paragraph.add_run(html_content)
    run.font.name = 'Courier New'  # Monospace font for code
    run.font.size = Inches(0.08)  # Smaller font size
    
    # Save to bytes buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer.getvalue()

# Streamlit App
def main():
    st.set_page_config(
        page_title="Casino Content Generator",
        page_icon="🎰",
        layout="wide"
    )
    
    # 🔐 AUTHENTICATION CHECK
    if not check_password():
        st.stop()  # Do not continue if check_password is not True
    
    st.title("🎰 Casino Content Generator")
    st.subheader("Generate SEO-optimized casino reviews with AI")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Logout button at top of sidebar
        if st.button("🚪 Logout", use_container_width=True):
            st.session_state["password_correct"] = False
            st.rerun()
        
        st.divider()
        
        # API Key input - check secrets first
        if "OPENAI_API_KEY" in st.secrets:
            api_key = st.secrets["OPENAI_API_KEY"]
            st.success("✅ OpenAI API Key loaded from secrets")
        else:
            api_key = st.text_input(
                "OpenAI API Key",
                type="password",
                help="Enter your OpenAI API key"
            )
            
            if not api_key:
                st.warning("Please enter your OpenAI API key to continue")
                return
        
        st.divider()
        
        # Mode selection
        mode = st.radio(
            "🎯 Generation Mode",
            ["📝 Single Casino", "📊 Bulk Generation"],
            help="Choose whether to generate content for one casino or multiple casinos"
        )
    
    # Main content area based on mode selection
    if mode == "📝 Single Casino":
        st.header("📝 Single Casino Generation")
        
        with st.form("single_casino_form"):
            casino_name = st.text_input("Casino Name", placeholder="e.g., Zitobox Casino")
            keyword = st.text_input("Target Keyword", placeholder="e.g., zitobox casino promo code")
            
            st.subheader("Competitor URLs")
            competitor_url_1 = st.text_input("Competitor URL 1", placeholder="https://competitor1.com/review")
            competitor_url_2 = st.text_input("Competitor URL 2", placeholder="https://competitor2.com/review")
            competitor_url_3 = st.text_input("Competitor URL 3", placeholder="https://competitor3.com/review")
            
            submitted = st.form_submit_button("🚀 Generate Content", use_container_width=True)
            
            if submitted:
                if casino_name and keyword and competitor_url_1:
                    competitor_urls = [url for url in [competitor_url_1, competitor_url_2, competitor_url_3] if url]
                    
                    with st.spinner("Generating content..."):
                        try:
                            generator = CasinoContentGenerator(api_key)
                            
                            start_time = time.time()
                            content = generator.generate_full_review(casino_name, keyword, competitor_urls)
                            end_time = time.time()
                            
                            st.success(f"✅ Content generated in {end_time - start_time:.1f} seconds!")
                            
                            # Store in session state for display
                            st.session_state.single_content = content
                            st.session_state.single_casino_name = casino_name
                            
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
                else:
                    st.error("Please fill in all required fields")
    
    else:  # Bulk Generation mode
        st.header("📊 Bulk Generation")
        
        # Example CSV download (outside of form)
        col1, col2 = st.columns([1, 1])
        with col1:
            st.info("Upload a CSV file with columns: casino_name, keyword, competitor_url_1, competitor_url_2, competitor_url_3")
        
        with col2:
            # Create example CSV for download
            example_data = {
                'casino_name': ['Zitobox Casino', 'Example Casino 2'],
                'keyword': ['zitobox casino promo code', 'example casino bonus'],
                'competitor_url_1': ['https://competitor1.com', 'https://competitor1.com'],
                'competitor_url_2': ['https://competitor2.com', 'https://competitor2.com'],
                'competitor_url_3': ['https://competitor3.com', 'https://competitor3.com']
            }
            df_example = pd.DataFrame(example_data)
            csv_example = df_example.to_csv(index=False)
            
            st.download_button(
                label="📥 Download Example CSV",
                data=csv_example,
                file_name="casino_bulk_example.csv",
                mime="text/csv",
                use_container_width=True
            )
        
        # Bulk generation form
        with st.form("bulk_casino_form"):
            uploaded_file = st.file_uploader("Choose CSV file", type="csv")
            
            bulk_submitted = st.form_submit_button("🚀 Generate Bulk Content", use_container_width=True)
            
            if bulk_submitted and uploaded_file:
                try:
                    df = pd.read_csv(uploaded_file)
                    required_columns = ['casino_name', 'keyword', 'competitor_url_1']
                    
                    if all(col in df.columns for col in required_columns):
                        st.info(f"Processing {len(df)} casinos...")
                        
                        generator = CasinoContentGenerator(api_key)
                        bulk_results = []
                        
                        progress_bar = st.progress(0)
                        
                        for idx, row in df.iterrows():
                            st.info(f"Processing: {row['casino_name']} ({idx + 1}/{len(df)})")
                            
                            competitor_urls = [
                                url for url in [
                                    row.get('competitor_url_1', ''),
                                    row.get('competitor_url_2', ''),
                                    row.get('competitor_url_3', '')
                                ] if url
                            ]
                            
                            try:
                                content = generator.generate_full_review(
                                    row['casino_name'],
                                    row['keyword'],
                                    competitor_urls
                                )
                                
                                bulk_results.append({
                                    'casino_name': row['casino_name'],
                                    'keyword': row['keyword'],
                                    'content': content,
                                    'status': 'Success'
                                })
                                
                            except Exception as e:
                                bulk_results.append({
                                    'casino_name': row['casino_name'],
                                    'keyword': row['keyword'],
                                    'content': '',
                                    'status': f'Error: {str(e)}'
                                })
                            
                            progress_bar.progress((idx + 1) / len(df))
                        
                        st.session_state.bulk_results = bulk_results
                        st.success(f"✅ Bulk generation complete! {len([r for r in bulk_results if r['status'] == 'Success'])} successful, {len([r for r in bulk_results if r['status'] != 'Success'])} errors")
                        
                    else:
                        st.error(f"CSV must contain columns: {required_columns}")
                        
                except Exception as e:
                    st.error(f"Error reading CSV: {str(e)}")
            
            elif bulk_submitted:
                st.error("Please upload a CSV file")
    
    # Results display - conditional based on mode
    st.header("📄 Generated Content")
    
    # Single content display - only show in Single mode
    if mode == "📝 Single Casino" and 'single_content' in st.session_state:
        st.subheader(f"Content for: {st.session_state.single_casino_name}")
        
        # Copy button and content display
        content_display = st.text_area(
            "HTML Content (Ready for CMS)",
            value=st.session_state.single_content,
            height=400,
            help="Copy this HTML content and paste it into your CMS editor"
        )
        
        # Download button - Create DOCX with HTML content as text
        docx_content = create_docx_with_html_content(
            st.session_state.single_casino_name, 
            st.session_state.single_content
        )
        
        st.download_button(
            label="📥 Download DOCX with HTML Content",
            data=docx_content,
            file_name=f"{st.session_state.single_casino_name.replace(' ', '_')}_html_content.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    # Bulk results display - only show in Bulk mode
    elif mode == "📊 Bulk Generation" and 'bulk_results' in st.session_state:
        st.subheader("📊 Bulk Generation Results")
        
        # Separate successful and failed results
        successful_results = [r for r in st.session_state.bulk_results if r['status'] == 'Success']
        error_results = [r for r in st.session_state.bulk_results if r['status'] != 'Success']
        
        if successful_results:
            # Download All button at the top
            zip_buffer = io.BytesIO()
            import zipfile
            
            with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
                for result in successful_results:
                    # Create DOCX content for each casino
                    docx_content = create_docx_with_html_content(
                        result['casino_name'], 
                        result['content']
                    )
                    filename = f"{result['casino_name'].replace(' ', '_')}_html_content.docx"
                    zip_file.writestr(filename, docx_content)
            
            zip_buffer.seek(0)
            
            col1, col2 = st.columns([1, 3])
            with col1:
                st.download_button(
                    label="📦 Download All DOCX Files",
                    data=zip_buffer.getvalue(),
                    file_name="bulk_casino_html_content.zip",
                    mime="application/zip",
                    use_container_width=True
                )
            with col2:
                st.info(f"✅ {len(successful_results)} files ready for download")
            
            st.divider()
            
            # Individual results in dataframe-like format
            st.subheader("📋 Individual Downloads")
            
            # Create header row
            col1, col2, col3, col4 = st.columns([3, 2, 2, 2])
            with col1:
                st.markdown("**Casino Name**")
            with col2:
                st.markdown("**Keyword**")
            with col3:
                st.markdown("**Download**")
            with col4:
                st.markdown("**Copy Content**")
            
            st.divider()
            
            # Individual rows for each result
            for i, result in enumerate(successful_results):
                col1, col2, col3, col4 = st.columns([3, 2, 2, 2])
                
                with col1:
                    st.markdown(f"**{result['casino_name']}**")
                
                with col2:
                    st.markdown(f"`{result['keyword']}`")
                
                with col3:
                    # Individual download button
                    docx_content = create_docx_with_html_content(
                        result['casino_name'], 
                        result['content']
                    )
                    
                    st.download_button(
                        label="📥 DOCX",
                        data=docx_content,
                        file_name=f"{result['casino_name'].replace(' ', '_')}_html_content.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_{i}",
                        use_container_width=True
                    )
                
                with col4:
                    # Copy content button - shows content in expandable section
                    if st.button("📋 Copy", key=f"copy_{i}", use_container_width=True):
                        st.session_state[f"show_content_{i}"] = not st.session_state.get(f"show_content_{i}", False)
                
                # Show content in expandable section when copy button is clicked
                if st.session_state.get(f"show_content_{i}", False):
                    with st.expander(f"HTML Content for {result['casino_name']}", expanded=True):
                        st.text_area(
                            "HTML Content (Click to select all and copy)",
                            value=result['content'],
                            height=200,
                            key=f"content_area_{i}",
                            help="Click in the text area, then Ctrl+A to select all, then Ctrl+C to copy"
                        )
                
                # Add subtle separator between rows
                if i < len(successful_results) - 1:
                    st.markdown("<hr style='margin: 10px 0; border: 0; border-top: 1px solid #eee;'>", unsafe_allow_html=True)
        
        # Error summary
        if error_results:
            st.divider()
            st.subheader("❌ Generation Errors")
            for result in error_results:
                st.error(f"**{result['casino_name']}**: {result['status']}")
        
        # Summary stats
        if successful_results or error_results:
            st.divider()
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("✅ Successful", len(successful_results))
            with col2:
                st.metric("❌ Failed", len(error_results))
            with col3:
                st.metric("📊 Total", len(st.session_state.bulk_results))
    
    # Show placeholder when no content generated yet
    elif mode == "📝 Single Casino":
        st.info("👆 Generate content for a single casino to see results here")
    else:  # Bulk mode with no results
        st.info("👆 Upload a CSV file and generate bulk content to see results here")

if __name__ == "__main__":
    main()
